import os
import subprocess
import logging
from backend.converters.base import BaseConverter
from backend.utils.sys_info import get_libreoffice_path

logger = logging.getLogger("office_converter")

class OfficeConverter(BaseConverter):
    def can_convert(self, from_ext: str, to_ext: str) -> bool:
        from_ext = from_ext.lower().strip('.')
        to_ext = to_ext.lower().strip('.')
        
        # Word docs can go to PDF, HTML, or TXT
        if from_ext in ["docx", "doc", "odt"]:
            return to_ext in ["pdf", "html", "txt"]
        return False

    def is_available(self) -> bool:
        # It's always "available" because it has a Python-only fallback for .docx
        return True

    def convert(self, input_path: str, output_path: str, options: dict = None) -> bool:
        from_ext = os.path.splitext(input_path)[1].lower().strip('.')
        to_ext = os.path.splitext(output_path)[1].lower().strip('.')

        # Check for LibreOffice first (best quality)
        lo_path = get_libreoffice_path()
        if lo_path:
            return self._convert_with_libreoffice(lo_path, input_path, output_path, to_ext)

        # Fallbacks when LibreOffice is not available
        logger.warning("LibreOffice not detected. Falling back to Python libraries.")
        
        if from_ext == "docx":
            if to_ext == "txt":
                return self._docx_to_txt(input_path, output_path)
            elif to_ext == "html":
                return self._docx_to_html(input_path, output_path)
            elif to_ext == "pdf":
                return self._docx_to_pdf_fallback(input_path, output_path)
        else:
            raise RuntimeError(
                f"Conversion from .{from_ext} to .{to_ext} without LibreOffice is not supported. "
                f"Please install LibreOffice for high-fidelity offline office document conversion."
            )

    def _convert_with_libreoffice(self, lo_path: str, input_path: str, output_path: str, to_ext: str) -> bool:
        logger.info(f"Converting using LibreOffice: {input_path} -> {to_ext}")
        out_dir = os.path.dirname(output_path)
        
        # LibreOffice headless command
        # For PDF, filter is 'writer_pdf_Export'
        # For HTML, filter is 'HTML'
        # For TXT, filter is 'Text (encoded)'
        filt = "writer_pdf_Export"
        if to_ext == "html":
            filt = "HTML"
        elif to_ext == "txt":
            filt = "Text (encoded)"

        cmd = [
            lo_path,
            "--headless",
            "--convert-to",
            f"{to_ext}:{filt}",
            "--outdir",
            out_dir,
            input_path
        ]
        
        try:
            # Run LibreOffice subprocess
            result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=30)
            if result.returncode != 0:
                logger.error(f"LibreOffice failed: {result.stderr}")
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

            # LibreOffice outputs to out_dir with the input file's name and new extension
            in_basename = os.path.basename(input_path)
            in_name_no_ext = os.path.splitext(in_basename)[0]
            lo_output_name = f"{in_name_no_ext}.{to_ext}"
            lo_output_path = os.path.join(out_dir, lo_output_name)

            # Move file to the exact requested output_path if different
            if os.path.exists(lo_output_path) and os.path.abspath(lo_output_path) != os.path.abspath(output_path):
                if os.path.exists(output_path):
                    os.remove(output_path)
                os.rename(lo_output_path, output_path)
                return True
            elif os.path.exists(output_path):
                return True
            else:
                raise FileNotFoundError(f"LibreOffice succeeded but output file was not found at {lo_output_path}")
        except subprocess.TimeoutExpired:
            logger.error("LibreOffice conversion timed out.")
            raise RuntimeError("LibreOffice conversion timed out.")

    def _docx_to_txt(self, input_path: str, output_path: str) -> bool:
        import docx
        doc = docx.Document(input_path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_txt = [cell.text for cell in row.cells]
                fullText.append(" | ".join(row_txt))
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(fullText))
        return True

    def _docx_to_html(self, input_path: str, output_path: str) -> bool:
        # We can use mammoth to get a clean HTML structure
        import mammoth
        with open(input_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            messages = result.messages
            for message in messages:
                logger.debug(f"Mammoth warning: {message.message}")
        
        # Wrap HTML in a simple page template
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8"/>
    <title>Converted Word Document</title>
    <style>
        body {{ font-family: sans-serif; line-height: 1.6; max-width: 800px; margin: 40px auto; padding: 0 20px; }}
        table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
        th, td {{ border: 1px solid #ccc; padding: 8px; text-align: left; }}
        th {{ background-color: #f4f4f4; }}
    </style>
</head>
<body>
    {html}
</body>
</html>"""
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_html)
        return True

    def _docx_to_pdf_fallback(self, input_path: str, output_path: str) -> bool:
        # Generates a simple PDF from DOCX paragraphs using ReportLab
        import docx
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors

        doc = docx.Document(input_path)
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Create custom Paragraph styles matching Word
        body_style = ParagraphStyle(
            'DocxBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=6
        )
        
        title_style = ParagraphStyle(
            'DocxTitle',
            parent=styles['Heading1'],
            fontSize=20,
            leading=24,
            spaceAfter=12
        )

        heading_style = ParagraphStyle(
            'DocxHeading',
            parent=styles['Heading2'],
            fontSize=14,
            leading=18,
            spaceAfter=8,
            spaceBefore=12
        )

        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                story.append(Spacer(1, 10))
                continue

            # Map Word style names to PDF paragraph styles
            p_style = body_style
            if "heading 1" in p.style.name.lower() or "title" in p.style.name.lower():
                p_style = title_style
            elif "heading" in p.style.name.lower():
                p_style = heading_style
            
            story.append(Paragraph(txt, p_style))

        # Add tables
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Renders cell text wrapped in a paragraph to support automatic height
                    row_data.append(Paragraph(cell.text, body_style))
                data.append(row_data)

            if data:
                t = Table(data, colWidths=[100]*len(data[0]))
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('GRID', (0,0), (-1,-1), 1, colors.grey),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                ]))
                story.append(Spacer(1, 15))
                story.append(t)
                story.append(Spacer(1, 15))

        pdf.build(story)
        return True
